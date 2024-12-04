from docx import Document
from docx.shared import Pt

doc = Document("hello_python.docx")

bold_text = []
for para in doc.paragraphs:
    for run in para.runs:
        if run.bold:
            bold_text.append(run.text)

print("Жирный текст в файле:", bold_text)

new_doc = Document()
paragraph = new_doc.add_paragraph()
run = paragraph.add_run("This is a paragraph with custom font and size.")

run.font.name = 'Arial'
run.font.size = Pt(16)

new_doc.save("custom_font_paragraph.docx")
